"""Tests for text extraction service."""

import os
import tempfile

from app.services.text_extraction import extract_text, extract_text_from_docx


def test_extract_text_from_docx():
    """Test DOCX text extraction using a real DOCX file."""
    from docx import Document as DocxDoc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = DocxDoc()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("This is a test document.")
        doc.save(f.name)
        path = f.name

    try:
        text = extract_text_from_docx(path)
        assert "Hello World" in text
        assert "test document" in text
    finally:
        os.unlink(path)


def test_extract_text_routes_by_extension():
    """Test that extract_text routes to the right extractor."""
    from docx import Document as DocxDoc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = DocxDoc()
        doc.add_paragraph("Routing test")
        doc.save(f.name)
        path = f.name

    try:
        text = extract_text(path)
        assert "Routing test" in text
    finally:
        os.unlink(path)


def test_extract_text_unsupported_extension():
    """Test that unsupported extensions return empty string."""
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"some content")
        path = f.name

    try:
        text = extract_text(path)
        assert text == ""
    finally:
        os.unlink(path)


def test_extract_text_nonexistent_file():
    """Test that nonexistent files return empty string."""
    text = extract_text("/nonexistent/file.docx")
    assert text == ""


def test_docx_table_extraction():
    """Test that text from tables is also extracted."""
    from docx import Document as DocxDoc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = DocxDoc()
        doc.add_paragraph("Header")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell A"
        table.cell(0, 1).text = "Cell B"
        table.cell(1, 0).text = "Cell C"
        table.cell(1, 1).text = "Cell D"
        doc.save(f.name)
        path = f.name

    try:
        text = extract_text_from_docx(path)
        assert "Header" in text
        assert "Cell A" in text
        assert "Cell D" in text
    finally:
        os.unlink(path)
