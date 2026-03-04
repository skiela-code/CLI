"""Tests for the DOCX template engine."""

import os
import tempfile

from docx import Document as DocxDocument

from app.services.template_engine import extract_placeholders, render_document


def _create_test_docx(content: str) -> str:
    """Helper to create a temp DOCX with given content."""
    doc = DocxDocument()
    doc.add_paragraph(content)
    path = os.path.join(tempfile.mkdtemp(), "test.docx")
    doc.save(path)
    return path


def test_extract_placeholders():
    path = _create_test_docx("Hello {{CLIENT_NAME}}, your deal {{DEAL_TITLE}} is ready.")
    tokens = extract_placeholders(path)
    assert "{{CLIENT_NAME}}" in tokens
    assert "{{DEAL_TITLE}}" in tokens
    assert len(tokens) == 2


def test_render_document():
    path = _create_test_docx("Agreement for {{CLIENT_NAME}} dated {{DATE}}.")
    out = render_document(path, {
        "{{CLIENT_NAME}}": "Acme Corp",
        "{{DATE}}": "2026-01-01",
    }, output_dir=tempfile.mkdtemp())

    assert os.path.exists(out)
    doc = DocxDocument(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Acme Corp" in text
    assert "2026-01-01" in text
    assert "{{CLIENT_NAME}}" not in text


def test_extract_no_placeholders():
    path = _create_test_docx("This document has no placeholders at all.")
    tokens = extract_placeholders(path)
    assert tokens == []
