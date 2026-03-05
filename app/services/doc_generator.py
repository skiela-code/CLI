"""Document generation service with async background jobs.

Architecture decision: We use asyncio.create_task() for background document
generation instead of Celery/Redis. Justification:
  1. V1 simplicity — no extra infrastructure (Redis, worker processes)
  2. FastAPI already runs in an async event loop
  3. Document generation is I/O-bound (API calls, file writes) — async is ideal
  4. We track job status in the DB, so jobs survive page refreshes
  5. Trade-off: jobs don't survive server restarts. Acceptable for V1.
"""

import asyncio
import os
import uuid
from datetime import datetime
from typing import Any, Optional

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.database import async_session_factory
from app.core.logging import log
from app.integrations.llm_router import get_llm_router
from app.models.models import (
    ContentBlock,
    Deal,
    DocStatus,
    DocType,
    Document,
    DocumentVersion,
    PricingLineItem,
    PricingTable,
    Template,
    TemplatePlaceholder,
    PlaceholderSource,
)
from app.services.red_flags import check_red_flags
from app.services.template_engine import render_document

settings = get_settings()

# Track running generation tasks
_running_tasks: dict[str, asyncio.Task] = {}


async def start_generation(
    db: AsyncSession,
    document_id: uuid.UUID,
    options: dict[str, Any] | None = None,
) -> None:
    """Launch document generation as a background asyncio task."""
    options = options or {}
    task = asyncio.create_task(_generate_document(str(document_id), options))
    _running_tasks[str(document_id)] = task
    log.info("generation_started", document_id=str(document_id))


async def _generate_document(document_id: str, options: dict[str, Any]) -> None:
    """Background task: resolve placeholders, call AI, render DOCX, run red-flag check."""
    async with async_session_factory() as db:
        try:
            doc = await db.get(Document, uuid.UUID(document_id))
            if not doc:
                log.error("document_not_found", document_id=document_id)
                return

            # Load related entities
            template = await db.get(Template, doc.template_id) if doc.template_id else None
            deal = await db.get(Deal, doc.deal_id) if doc.deal_id else None

            # Load pricing
            pricing_items: list[dict] = []
            if doc.pricing_table_id:
                pt = await db.get(PricingTable, doc.pricing_table_id)
                if pt:
                    items_q = await db.execute(
                        select(PricingLineItem)
                        .where(PricingLineItem.pricing_table_id == pt.id)
                        .order_by(PricingLineItem.sort_order)
                    )
                    for item in items_q.scalars():
                        pricing_items.append({
                            "description": item.description,
                            "quantity": item.quantity,
                            "unit_price": item.unit_price,
                            "discount_pct": item.discount_pct,
                            "total": item.quantity * item.unit_price * (1 - item.discount_pct / 100),
                        })

            # Build replacements
            replacements: dict[str, str] = {}
            generated_sections: dict[str, str] = {}

            if template:
                placeholders_q = await db.execute(
                    select(TemplatePlaceholder).where(TemplatePlaceholder.template_id == template.id)
                )
                router = get_llm_router()
                deal_ctx = {
                    "title": deal.title if deal else "N/A",
                    "org_name": deal.org_name if deal else "N/A",
                    "contact_name": deal.contact_name if deal else "N/A",
                    "contact_email": deal.contact_email if deal else "",
                    "value": deal.value if deal else 0,
                }

                for ph in placeholders_q.scalars():
                    value = ph.default_value or ""

                    if ph.source == PlaceholderSource.DEAL_FIELD and deal:
                        field = ph.source_field or ""
                        if field.startswith("deal."):
                            attr = field.split(".", 1)[1]
                            value = str(getattr(deal, attr, "")) or value
                        elif field.startswith("custom."):
                            key = field.split(".", 1)[1]
                            value = str((deal.custom_fields or {}).get(key, "")) or value

                    elif ph.source == PlaceholderSource.CLIENT_FIELD and deal:
                        field = ph.source_field or ""
                        if "org" in field:
                            value = deal.org_name or value
                        elif "contact" in field or "name" in field:
                            value = deal.contact_name or value
                        elif "email" in field:
                            value = deal.contact_email or value

                    elif ph.source == PlaceholderSource.CONTENT_BLOCK and ph.content_block_id:
                        block = await db.get(ContentBlock, ph.content_block_id)
                        if block:
                            value = block.body

                    elif ph.source == PlaceholderSource.AI_GENERATED:
                        prompt = ph.ai_prompt or f"Generate content for: {ph.label or ph.token}"
                        system = "You are a document section writer. Generate professional content."
                        user_msg = f"{prompt}\n\nContext:\n"
                        for ctx_k, ctx_v in deal_ctx.items():
                            user_msg += f"  {ctx_k}: {ctx_v}\n"
                        value = await router.generate(system, user_msg)
                        generated_sections[ph.token] = value

                    replacements[ph.token] = value

                # Handle pricing table placeholder
                if "{{PRICING_TABLE}}" in replacements or any("PRICING" in k for k in replacements):
                    pricing_text = _format_pricing_text(pricing_items)
                    replacements["{{PRICING_TABLE}}"] = pricing_text

                # For commercial offers, generate narrative
                length = options.get("length", "medium")
                if doc.doc_type == DocType.COMMERCIAL_OFFER or doc.doc_type == "commercial_offer":
                    narrative_system = (
                        "You are a professional proposal writer. Generate clear, persuasive "
                        "business document sections. Use the pricing data as the single source of truth. "
                        "Do NOT invent prices or quantities — reference them from the provided table."
                    )
                    narrative_user = (
                        f"Document type: commercial_offer\n"
                        f"Length: {length}\n"
                        f"Deal: {deal_ctx.get('title', 'N/A')}\n"
                        f"Client: {deal_ctx.get('org_name', 'N/A')}\n"
                        f"Contact: {deal_ctx.get('contact_name', 'N/A')}\n\n"
                        f"Pricing table:\n"
                    )
                    for p_item in pricing_items:
                        narrative_user += (
                            f"  - {p_item['description']}: {p_item['quantity']}x @ {p_item['unit_price']} "
                            f"(discount {p_item.get('discount_pct', 0)}%)\n"
                        )
                    narrative_user += (
                        "\nWrite a professional narrative section that:\n"
                        "1. Introduces the solution\n"
                        "2. References the pricing accurately\n"
                        "3. Highlights value propositions\n"
                        "4. Includes next steps\n"
                    )
                    narrative = await router.generate(narrative_system, narrative_user)
                    generated_sections["narrative"] = narrative
                    # Fill AI placeholders if not already filled
                    for key in ["{{AI_EXECUTIVE_SUMMARY}}", "{{AI_SOLUTION_DESCRIPTION}}",
                                "{{AI_VALUE_PROPOSITION}}", "{{AI_NEXT_STEPS}}"]:
                        if key not in replacements or not replacements[key]:
                            replacements[key] = narrative

            # Fill any remaining date placeholders
            from datetime import date
            today = date.today().isoformat()
            replacements.setdefault("{{DATE}}", today)

            # Render DOCX
            if template and template.file_path and os.path.exists(template.file_path):
                out_path = render_document(template.file_path, replacements)
            else:
                # Create a basic doc with content
                from docx import Document as DocxDoc
                d = DocxDoc()
                d.add_heading(doc.title, level=1)
                for section_name, text in replacements.items():
                    d.add_heading(section_name.strip("{} "), level=2)
                    d.add_paragraph(text)
                if generated_sections.get("narrative"):
                    d.add_heading("Narrative", level=2)
                    for line in generated_sections["narrative"].split("\n"):
                        d.add_paragraph(line)

                os.makedirs(settings.generated_docs_path, exist_ok=True)
                out_path = os.path.join(settings.generated_docs_path, f"{uuid.uuid4().hex}.docx")
                d.save(out_path)

            # Red flags check
            rf_result = check_red_flags(
                content=replacements if replacements else {"body": ""},
                deal={"org_name": deal.org_name, "contact_name": deal.contact_name} if deal else None,
                pricing=pricing_items if pricing_items else None,
            )

            # Determine version number
            ver_q = await db.execute(
                select(func.coalesce(func.max(DocumentVersion.version_number), 0))
                .where(DocumentVersion.document_id == doc.id)
            )
            next_version = ver_q.scalar() + 1

            # Save version
            version = DocumentVersion(
                document_id=doc.id,
                version_number=next_version,
                file_path=out_path,
                generated_content={
                    "replacements": {k: v[:500] for k, v in replacements.items()},
                    "ai_sections": list(generated_sections.keys()),
                    "options": options,
                },
            )
            db.add(version)

            # Update document
            doc.red_flags = rf_result
            doc.status = DocStatus.DRAFT

            await db.commit()
            log.info("generation_complete", document_id=document_id, version=next_version)

        except Exception as e:
            log.error("generation_failed", document_id=document_id, error=str(e))
            await db.rollback()
            raise
        finally:
            _running_tasks.pop(document_id, None)


def _format_pricing_text(items: list[dict]) -> str:
    """Format pricing items as readable text for DOCX insertion."""
    if not items:
        return "No pricing data available."

    lines = ["Item | Qty | Unit Price | Discount | Total", "-" * 60]
    grand_total = 0
    for item in items:
        total = item.get("total", item["quantity"] * item["unit_price"])
        grand_total += total
        lines.append(
            f"{item['description']} | {item['quantity']} | "
            f"{item['unit_price']:.2f} | {item.get('discount_pct', 0)}% | {total:.2f}"
        )
    lines.append("-" * 60)
    lines.append(f"TOTAL: {grand_total:.2f}")
    return "\n".join(lines)


def get_task_status(document_id: str) -> str:
    """Check if a generation task is still running."""
    task = _running_tasks.get(document_id)
    if task is None:
        return "completed"
    if task.done():
        _running_tasks.pop(document_id, None)
        if task.exception():
            return "failed"
        return "completed"
    return "running"
