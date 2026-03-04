"""DOCX template engine: extract placeholders and render documents."""

import os
import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

from app.core.config import get_settings
from app.core.logging import log

settings = get_settings()

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z_][A-Z0-9_]*)\}\}")


def extract_placeholders(file_path: str) -> list[str]:
    """Extract {{PLACEHOLDER}} tokens from a DOCX file."""
    doc = DocxDocument(file_path)
    tokens: set[str] = set()

    for para in doc.paragraphs:
        # Reconstruct full text from runs to handle split placeholders
        full_text = "".join(run.text for run in para.runs)
        for match in PLACEHOLDER_RE.finditer(full_text):
            tokens.add(match.group(0))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = "".join(run.text for run in para.runs)
                    for match in PLACEHOLDER_RE.finditer(full_text):
                        tokens.add(match.group(0))

    return sorted(tokens)


def render_document(
    template_path: str,
    replacements: dict[str, str],
    output_dir: str | None = None,
) -> str:
    """Replace placeholders in DOCX template and save to output path. Returns output file path."""
    doc = DocxDocument(template_path)

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text
        for token, value in replacements.items():
            new_text = new_text.replace(token, value)
        if new_text != full_text:
            # Clear all runs and put new text in first run
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = new_text
                else:
                    run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    out_dir = output_dir or settings.generated_docs_path
    os.makedirs(out_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)
    log.info("document_rendered", output=out_path)
    return out_path


def create_sample_template(doc_type: str = "contract") -> str:
    """Create a sample DOCX template with common placeholders. Returns file path."""
    doc = DocxDocument()

    if doc_type == "nda":
        doc.add_heading("Non-Disclosure Agreement", level=1)
        doc.add_paragraph(
            "This Non-Disclosure Agreement (\"Agreement\") is entered into on {{DATE}} "
            "by and between:"
        )
        doc.add_paragraph("Party A: {{COMPANY_NAME}}, represented by {{COMPANY_REPRESENTATIVE}}")
        doc.add_paragraph("Party B: {{CLIENT_NAME}}, represented by {{CLIENT_REPRESENTATIVE}}")
        doc.add_paragraph("")
        doc.add_heading("1. Confidential Information", level=2)
        doc.add_paragraph("{{CONFIDENTIALITY_SCOPE}}")
        doc.add_heading("2. Obligations", level=2)
        doc.add_paragraph("{{OBLIGATIONS_TEXT}}")
        doc.add_heading("3. Term", level=2)
        doc.add_paragraph(
            "This Agreement shall remain in effect for {{DURATION}} from the date of signing."
        )
        doc.add_heading("4. Governing Law", level=2)
        doc.add_paragraph("This Agreement shall be governed by the laws of {{JURISDICTION}}.")
    elif doc_type == "purchase_annex":
        doc.add_heading("Purchase Annex", level=1)
        doc.add_paragraph("Annex to Contract No. {{CONTRACT_NUMBER}}")
        doc.add_paragraph("Date: {{DATE}}")
        doc.add_paragraph("Supplier: {{COMPANY_NAME}}")
        doc.add_paragraph("Buyer: {{CLIENT_NAME}}")
        doc.add_heading("Items", level=2)
        doc.add_paragraph("{{PRICING_TABLE}}")
        doc.add_heading("Delivery Terms", level=2)
        doc.add_paragraph("{{DELIVERY_TERMS}}")
        doc.add_heading("Payment Terms", level=2)
        doc.add_paragraph("{{PAYMENT_TERMS}}")
    elif doc_type == "commercial_offer":
        doc.add_heading("Commercial Offer", level=1)
        doc.add_paragraph("Prepared for: {{CLIENT_NAME}}")
        doc.add_paragraph("Date: {{DATE}}")
        doc.add_paragraph("Reference: {{DEAL_TITLE}}")
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph("{{AI_EXECUTIVE_SUMMARY}}")
        doc.add_heading("Proposed Solution", level=2)
        doc.add_paragraph("{{AI_SOLUTION_DESCRIPTION}}")
        doc.add_heading("Investment", level=2)
        doc.add_paragraph("{{PRICING_TABLE}}")
        doc.add_heading("Value Proposition", level=2)
        doc.add_paragraph("{{AI_VALUE_PROPOSITION}}")
        doc.add_heading("Next Steps", level=2)
        doc.add_paragraph("{{AI_NEXT_STEPS}}")
    else:  # contract
        doc.add_heading("Service Agreement", level=1)
        doc.add_paragraph("Contract No: {{CONTRACT_NUMBER}}")
        doc.add_paragraph("Date: {{DATE}}")
        doc.add_paragraph("")
        doc.add_paragraph("This Agreement is between {{COMPANY_NAME}} (\"Provider\") and {{CLIENT_NAME}} (\"Client\").")
        doc.add_heading("1. Scope of Services", level=2)
        doc.add_paragraph("{{SCOPE_OF_SERVICES}}")
        doc.add_heading("2. Term", level=2)
        doc.add_paragraph("Start: {{START_DATE}} — End: {{END_DATE}}")
        doc.add_heading("3. Compensation", level=2)
        doc.add_paragraph("{{PRICING_TABLE}}")
        doc.add_heading("4. Terms & Conditions", level=2)
        doc.add_paragraph("{{TERMS_AND_CONDITIONS}}")
        doc.add_heading("5. Signatures", level=2)
        doc.add_paragraph("Provider: {{COMPANY_REPRESENTATIVE}}    Client: {{CLIENT_REPRESENTATIVE}}")

    out_dir = settings.upload_path
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, f"sample_{doc_type}.docx")
    doc.save(path)
    return path
