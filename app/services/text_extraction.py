"""Text extraction from DOCX and PDF files."""

import os

from app.core.logging import log


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    from docx import Document as DocxDocument
    try:
        doc = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        return "\n".join(parts)
    except Exception as e:
        log.error("docx_extraction_failed", file=file_path, error=str(e))
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        return "\n".join(parts)
    except Exception as e:
        log.error("pdf_extraction_failed", file=file_path, error=str(e))
        return ""


def extract_text(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        log.warning("unsupported_file_type", file=file_path, ext=ext)
        return ""
